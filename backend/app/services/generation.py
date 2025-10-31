import re
import yaml
from os import path
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from docx.oxml.parser import OxmlElement
from openai import OpenAI
from ..core.config import settings
from ..schemas.document import GenerateRequest

# Определяем абсолютный путь к директории с шаблонами
# Это делает путь независимым от того, откуда запускается приложение
BASE_DIR = path.dirname(path.dirname(path.abspath(__file__)))
TEMPLATES_DIR = path.join(BASE_DIR, '..', 'templates')

def generate_yaml_from_gpt(master_prompt: str, user_data: GenerateRequest) -> str | None:
    client = OpenAI(
        api_key=settings.BOT_HUB_TOKEN,
        base_url=settings.BOT_HUB_PROXY_URL
    )
    
    user_prompt_text = user_data.user_prompt
    user_data_json = user_data.model_dump_json(exclude={'user_prompt'}) # Exclude from JSON to avoid duplication

    # Construct a clearer message for the AI
    user_content = (
        f"Вот текстовый запрос от пользователя:\n--- НАЧАЛО ЗАПРОСА ---\n{user_prompt_text}\n--- КОНЕЦ ЗАПРОСА ---\n\n"
        f"А вот остальные структурированные данные из формы:\n{user_data_json}\n\n"
        "Твоя задача — объединить всю эту информацию и сгенерировать YAML строго по правилам из системного промпта."
    )

    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    'role': 'system',
                    'content': master_prompt
                },
                {
                    'role': 'user',
                    'content': user_content
                }
            ],
            model='yandexgpt',#yandexgpt-lite
        )
        
        result_text = str(chat_completion.choices[0].message.content)
        cleaned_yaml = result_text.strip().replace("```yaml", "").replace("```", "").strip()
        
        return cleaned_yaml
    except Exception as e:
        print(f"Error calling BotHub API: {e}")
        return None

def fill_docx_template(template_name: str, data: dict) -> BytesIO:
    template_path = path.join(TEMPLATES_DIR, template_name)
    try:
        doc = Document(template_path)
    except Exception:
        raise FileNotFoundError(f"Template not found at {template_path}")

    # Uppercase the document_type before creating replacements
    if 'document_type' in data and data['document_type']:
        data['document_type'] = str(data['document_type']).upper()

    # Create a flat dictionary of all possible tags for replacement
    replacements = {}
    for key, value in data.items():
        if isinstance(value, dict):
            for sub_key, sub_value in value.items():
                # Replace None with an empty string, otherwise convert to string
                val = str(sub_value) if sub_value is not None else ""
                replacements[f"{{{{ {key}.{sub_key} }}}}"] = val
                replacements[f"{{{{{key}.{sub_key}}}}}"] = val
        else:
            # Replace None with an empty string, otherwise convert to string
            val = str(value) if value is not None else ""
            replacements[f"{{{{ {key} }}}}"] = val
            replacements[f"{{{{{key}}}}}"] = val

    # Function to handle multi-paragraph replacement for the 'body'
    def replace_body_text(doc, body_text):
        p_to_replace = None
        for p in doc.paragraphs:
            if '{{ body }}' in p.text or '{{body}}' in p.text:
                p_to_replace = p
                break
        
        if p_to_replace:
            lines = [line.strip() for line in body_text.split('\n') if line.strip()]
            
            # Set the first line of text in the placeholder paragraph
            p_to_replace.text = lines[0] if lines else ""
            p_to_replace.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_to_replace.paragraph_format.first_line_indent = Cm(1.25)
            
            # Store the style from the now-filled placeholder paragraph
            original_style = p_to_replace.style
            
            # Start adding new paragraphs after the placeholder
            p_element = p_to_replace._p # The <w:p> XML element
            for line in lines[1:]:
                # Create a new <w:p> element and add it after the previous one
                new_p_element = OxmlElement("w:p")
                p_element.addnext(new_p_element)
                
                # Create a Paragraph object from the new XML element
                new_p = Paragraph(new_p_element, p_to_replace._parent)
                new_p.text = line
                
                # Apply the same style and alignment
                new_p.style = original_style
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.first_line_indent = Cm(1.25)
                
                # Update the reference to the last element for the next iteration
                p_element = new_p_element

    # A robust search and replace function that handles bolding for header fields
    def search_and_replace(doc_part, replacements):
        for p in doc_part.paragraphs:
            original_text = p.text
            
            # Check if this paragraph contains a header tag that needs bolding
            # We do this before replacement
            needs_bolding = False
            for tag in replacements:
                if "_in_" in tag and tag in original_text:
                    needs_bolding = True
                    break

            # Perform all replacements in the text string
            modified_text = original_text
            for tag, value in replacements.items():
                modified_text = modified_text.replace(tag, str(value))
            
            # If any changes were made, update the paragraph's text.
            # Assigning to .text clears all runs and adds a new one.
            if original_text != modified_text:
                p.text = modified_text
                # Apply bolding to the entire paragraph if it was marked
                if needs_bolding:
                    for run in p.runs:
                        run.bold = True

        for table in doc_part.tables:
            for row in table.rows:
                for cell in row.cells:
                    search_and_replace(cell, replacements)

    # Create a dictionary of replacements, excluding the 'body'
    simple_replacements = {k: v for k, v in replacements.items() if 'body' not in k}

    # Replace all simple tags first, across the entire document
    search_and_replace(doc, simple_replacements)
    for section in doc.sections:
        search_and_replace(section.header, simple_replacements)
        search_and_replace(section.footer, simple_replacements)

    # Handle multi-paragraph body replacement last
    if 'body' in data and data['body']:
        replace_body_text(doc, str(data['body']))

    # Final, more robust cleanup pass to remove any remaining, unfilled tags
    tag_regex = re.compile(r'\{\{.*?\}\}')
    for p in doc.paragraphs:
        if tag_regex.search(p.text):
            # This is a more aggressive replacement that might lose some formatting in rare cases,
            # but it's much more reliable for removing leftover tags.
            p.text = tag_regex.sub('', p.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if tag_regex.search(p.text):
                        p.text = tag_regex.sub('', p.text)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
